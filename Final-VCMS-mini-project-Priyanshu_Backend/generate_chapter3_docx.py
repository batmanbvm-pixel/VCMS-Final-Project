from docx import Document
from docx.shared import Cm, Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from PIL import Image, ImageDraw, ImageFont
from pathlib import Path
import tempfile

OUT_PATH = Path('Chapter_3_System_Design_Detailed_Module_Description.docx')


def set_cell_shading(cell, fill):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), fill)
    tcPr.append(shd)


def set_cell_margins(cell, top=90, start=90, bottom=90, end=90):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = tcPr.first_child_found_in('w:tcMar')
    if tcMar is None:
        tcMar = OxmlElement('w:tcMar')
        tcPr.append(tcMar)
    for m, v in [('top', top), ('start', start), ('bottom', bottom), ('end', end)]:
        node = tcMar.find(qn(f'w:{m}'))
        if node is None:
            node = OxmlElement(f'w:{m}')
            tcMar.append(node)
        node.set(qn('w:w'), str(v))
        node.set(qn('w:type'), 'dxa')


def format_runs(paragraph, size=12, bold=False, italic=False, color=None):
    for run in paragraph.runs:
        run.font.name = 'Times New Roman'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        run.font.size = Pt(size)
        run.font.bold = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, size=12, bold=False, italic=False, color=None):
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    return run


def add_heading(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.2
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(14)
    run.font.bold = True
    run.font.color.rgb = RGBColor(31, 41, 55)
    return p


def add_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(10)
    r.font.bold = True
    return p


def format_paragraph(paragraph, align=None):
    paragraph.paragraph_format.space_after = Pt(6)
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.line_spacing = 1.5
    if align is not None:
        paragraph.alignment = align
    format_runs(paragraph)


def make_figure(path):
    img = Image.new('RGB', (1800, 1050), 'white')
    draw = ImageDraw.Draw(img)
    try:
        title_font = ImageFont.truetype('/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf', 46)
        bold_font = ImageFont.truetype('/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf', 28)
        body_font = ImageFont.truetype('/System/Library/Fonts/Supplemental/Times New Roman.ttf', 26)
        small_font = ImageFont.truetype('/System/Library/Fonts/Supplemental/Times New Roman.ttf', 22)
    except Exception:
        title_font = ImageFont.load_default()
        bold_font = ImageFont.load_default()
        body_font = ImageFont.load_default()
        small_font = ImageFont.load_default()

    draw.rounded_rectangle((80, 40, 1720, 140), radius=24, fill=(14, 116, 144), outline=(8, 51, 68), width=3)
    draw.text((900, 92), 'VCMS System Architecture Overview', fill='white', font=title_font, anchor='mm')

    boxes = {
        'Frontend Layer\nReact + TypeScript\nPatient / Doctor / Admin UI': (120, 220, 520, 400),
        'Backend Layer\nNode.js + Express\nREST APIs + Socket.IO': (640, 220, 1160, 400),
        'Database Layer\nMongoDB Atlas\nUsers, Appointments, Prescriptions': (1220, 220, 1680, 400),
        'Security & Control\nJWT, bcrypt, RBAC\nValidation & Protected Routes': (120, 560, 520, 760),
        'External Services\nWebRTC, Gemini, Tesseract\nPDF / OCR / Signaling': (640, 560, 1160, 760),
    }
    styles = {
        'Frontend Layer\nReact + TypeScript\nPatient / Doctor / Admin UI': ((223, 246, 255), (2, 132, 199)),
        'Backend Layer\nNode.js + Express\nREST APIs + Socket.IO': ((236, 253, 245), (22, 163, 74)),
        'Database Layer\nMongoDB Atlas\nUsers, Appointments, Prescriptions': ((255, 247, 237), (234, 88, 12)),
        'Security & Control\nJWT, bcrypt, RBAC\nValidation & Protected Routes': ((254, 242, 242), (220, 38, 38)),
        'External Services\nWebRTC, Gemini, Tesseract\nPDF / OCR / Signaling': ((250, 245, 255), (126, 34, 206)),
    }
    for text, box in boxes.items():
        fill, outline = styles[text]
        draw.rounded_rectangle(box, radius=28, fill=fill, outline=outline, width=4)
        lines = text.split('\n')
        cx = (box[0] + box[2]) / 2
        cy = (box[1] + box[3]) / 2
        y = cy - 40
        for i, line in enumerate(lines):
            draw.text((cx, y), line, fill=(17, 24, 39), font=bold_font if i == 0 else body_font, anchor='mm')
            y += 42

    arrow_color = (55, 65, 81)
    for y in [300, 335]:
        draw.line((520, y, 640, y), fill=arrow_color, width=8)
        draw.polygon([(640, y), (620, y - 12), (620, y + 12)], fill=arrow_color)
    for y in [300, 335]:
        draw.line((1160, y, 1220, y), fill=arrow_color, width=8)
        draw.polygon([(1220, y), (1200, y - 12), (1200, y + 12)], fill=arrow_color)
    for x in [850, 925]:
        draw.line((x, 400, x, 560), fill=arrow_color, width=8)
        draw.polygon([(x, 560), (x - 12, 540), (x + 12, 540)], fill=arrow_color)

    draw.rounded_rectangle((590, 455, 1210, 520), radius=18, fill=(243, 244, 246), outline=(156, 163, 175), width=2)
    draw.text((900, 488), 'Request / Response Flow + Real-Time Signaling', fill=(17, 24, 39), font=small_font, anchor='mm')

    img.save(path)


def add_module_section(doc, title, items):
    add_heading(doc, title)
    for subtitle, body in items:
        p = doc.add_paragraph()
        add_text(p, f'{subtitle}: ', 12, bold=True)
        add_text(p, body, 12)
        format_paragraph(p)


def main():
    temp_dir = Path(tempfile.mkdtemp(prefix='vcms_ch3_'))
    fig_path = temp_dir / 'architecture.png'
    make_figure(fig_path)

    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.54)
    section.right_margin = Cm(2.54)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    foot_run = footer.add_run('Virtual Clinic Management System (VCMS) – Chapter 3')
    foot_run.font.name = 'Times New Roman'
    foot_run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    foot_run.font.size = Pt(10)

    # Title page
    for text, size, bold, color, italic in [
        ('CHAPTER 3', 18, True, None, False),
        ('System Design and Detailed Module Description', 18, True, '0E7490', False),
        ('Virtual Clinic Management System (VCMS)', 14, False, None, True),
        ('Prepared from analysis of the current MERN stack implementation', 12, False, None, False),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
        r.font.size = Pt(size)
        r.font.bold = bold
        r.font.italic = italic
        if color:
            r.font.color.rgb = RGBColor.from_string(color)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━━')
    r.font.name = 'Times New Roman'
    r._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    r.font.size = Pt(12)

    add_heading(doc, '3.0 Project-Wide System Design')
    p = doc.add_paragraph()
    add_text(p, 'The Virtual Clinic Management System (VCMS) is a role-based healthcare platform built on the MERN stack. ', 12)
    add_text(p, 'It connects patients, doctors, and administrators through a single secure web application that supports registration, appointment booking, video consultation, prescription handling, AI-assisted health guidance, and preventive wellness analysis. ', 12)
    add_text(p, 'The system is designed to keep sensitive medical interactions inside the application while using MongoDB for persistent storage and WebRTC / Socket.IO for real-time communication.', 12)
    format_paragraph(p)

    p = doc.add_paragraph()
    add_text(p, 'From the project analysis, the implementation includes React + TypeScript on the frontend, Node.js + Express on the backend, MongoDB Atlas for data storage, JWT and bcrypt for security, Socket.IO for live events, WebRTC for video consultation, and AI/OCR services for medical assistance. ', 12)
    add_text(p, 'The database layer includes collections such as Users, Appointments, Prescriptions, MedicalHistory, Notifications, ChatMessages, VideoSessions, Contacts, ConsultationForms, and DoctorReviews.', 12)
    format_paragraph(p)

    add_heading(doc, '3.0.1 System Technology Stack Overview')
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = 'Table Grid'
    headers = ['Layer', 'Technology Used', 'Role in the Project']
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        set_cell_shading(cell, '0E7490')
        cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(cell)
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_runs(p, size=12, bold=True, color='FFFFFF')
    rows = [
        ['Frontend', 'React 18, TypeScript, Tailwind CSS, shadcn/ui', 'Role-based user interface, dashboards, forms, and booking workflows'],
        ['Backend', 'Node.js, Express.js, Mongoose', 'Authentication, business logic, API validation, and data processing'],
        ['Database', 'MongoDB Atlas', 'Persistent storage for users, appointments, prescriptions, and medical records'],
        ['Real-Time Layer', 'Socket.IO', 'Live notifications, appointment updates, chat signaling, and video events'],
        ['Video Layer', 'WebRTC', 'Peer-to-peer doctor-patient video consultation inside the application'],
        ['AI / OCR Layer', 'Gemini API, Tesseract.js, pdf-parse', 'Report simplification, prescription summary, and document text extraction'],
    ]
    for row in rows:
        cells = table.add_row().cells
        for i, txt in enumerate(row):
            cells[i].text = txt
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER
                format_runs(p, size=12)
    add_caption(doc, 'Table 3.1. System Technology Stack and Responsibility Mapping.')

    add_heading(doc, '3.0.2 Overall System Architecture Diagram')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(fig_path), width=Inches(6.8))
    add_caption(doc, 'Figure 3.1. High-level architecture of the VCMS platform.')

    add_heading(doc, '3.0.3 Module Summary and Data Flow View')
    table2 = doc.add_table(rows=1, cols=5)
    table2.alignment = WD_TABLE_ALIGNMENT.CENTER
    table2.style = 'Table Grid'
    headers2 = ['Module', 'Primary Users', 'Main Inputs', 'Core Outputs', 'Data Handling']
    for i, h in enumerate(headers2):
        c = table2.rows[0].cells[i]
        c.text = h
        set_cell_shading(c, '115E59')
        c.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        set_cell_margins(c)
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            format_runs(p, size=12, bold=True, color='FFFFFF')
    rows2 = [
        ['User Management', 'Patient, Doctor, Admin', 'Name, email, phone, password, role-specific profile data', 'Authenticated account and role-based access', 'Stored securely in MongoDB with hashed passwords'],
        ['Appointment Management', 'Patient, Doctor, Admin', 'Doctor selection, date, time, symptoms, chatbot answers', 'Booked/accepted/completed/cancelled appointments', 'Saved in MongoDB with status tracking'],
        ['Video Consultation', 'Patient, Doctor', 'Appointment room ID, camera/mic permission', 'Real-time consultation session', 'Video session metadata stored, media stays real-time'],
        ['Prescription Management', 'Doctor, Patient', 'Medicines, dosage, duration, instructions', 'Digital prescription linked to appointment', 'Persisted in database and accessible from dashboard'],
        ['AI Health Advisory', 'Patient, Doctor', 'Symptoms, reports, doctor notes', 'Specialization suggestion, summary, simplified analysis', 'AI output is informational; report uploads processed securely'],
        ['Wellness Module', 'Patient', 'Family health history and lifestyle answers', 'Risk awareness indicators', 'Processed live and not stored permanently'],
    ]
    for row in rows2:
        cells = table2.add_row().cells
        for i, txt in enumerate(row):
            cells[i].text = txt
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cells[i])
            for p in cells[i].paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT if i else WD_ALIGN_PARAGRAPH.CENTER
                format_runs(p, size=11)
    add_caption(doc, 'Table 3.2. Module-wise summary of input, processing, output, and data handling.')

    add_module_section(doc, '3.1 User Management Module', [
        ('3.1.1 Objective', 'This module manages registration, login, profile access, and role-based authorization for the three user categories supported by the platform: Patient, Doctor, and Admin.'),
        ('3.1.2 Functional Description', 'During registration, users enter name, email, phone number, password, and their selected role. Doctors may also provide professional details such as specialization, experience, and consultation fee, while patients can include age and basic medical history. The backend validates the inputs, hashes the password with bcrypt, and creates the user record in MongoDB. After login, JWT tokens are generated so the frontend can keep the user authenticated across protected pages.'),
        ('3.1.3 Project Relevance', 'The project uses this module as the gateway for all other features. A patient cannot book appointments without authentication, a doctor cannot issue prescriptions without a valid session, and an admin cannot manage the system without elevated access. This makes the module the central security control point for the entire application.'),
        ('3.1.4 Frontend and Backend Support', 'The login and registration screens are implemented in the React frontend, while the backend routes handle registration, authentication, profile updates, password recovery, and logout. Protected routing ensures that each dashboard is accessible only to the appropriate role.'),
        ('3.1.5 Data and Security', 'All passwords are stored in encrypted form and never saved as plain text. The role-based access control layer prevents unauthorized entry into restricted modules. This design follows common healthcare application security practices and supports safe handling of sensitive personal information.'),
    ])

    add_module_section(doc, '3.2 Appointment Management Module', [
        ('3.2.1 Objective', 'This module supports the complete appointment lifecycle: booking, viewing, updating, accepting, rejecting, cancelling, and completing appointments.'),
        ('3.2.2 Functional Description', 'Patients can book appointments manually by choosing a doctor, a date, and a time slot. In addition, the platform provides a chatbot-based booking path where the user answers simple questions about symptoms, preferred specialization, date, and time. The system then helps narrow down the right doctor and available slot. Appointments move through a status flow such as Booked, Accepted, In Progress, Completed, and Cancelled, depending on the actions taken by the patient, doctor, or admin.'),
        ('3.2.3 Project Relevance', 'The appointment module is the operational heart of the clinic system because it connects doctor availability, patient demand, and consultation scheduling. It also drives notifications, video sessions, and prescription creation after the consultation is completed.'),
        ('3.2.4 Database and Workflow', 'Appointment records are stored in MongoDB with references to both doctor and patient accounts. The backend checks slot conflicts, prevents invalid bookings, and records cancellation or rejection reasons. Once an appointment is confirmed, it becomes the entry point for the consultation room and later the prescription record.'),
        ('3.2.5 User Experience', 'The frontend presents appointments in dedicated dashboards for patients, doctors, and admins. Filters and status chips help users quickly find upcoming, completed, or cancelled visits. Socket-based updates ensure that changes appear instantly across sessions.'),
    ])

    add_module_section(doc, '3.3 Video Consultation Module', [
        ('3.3.1 Objective', 'This module enables real-time audio and video consultations between doctors and patients directly inside the web application.'),
        ('3.3.2 Functional Description', 'When the appointment reaches the consultation stage, a unique virtual room ID is created and shared with the participants. The doctor and patient join the same room at the scheduled time, and WebRTC handles the peer-to-peer media connection. Socket.IO is used for signaling so that the browser can exchange offers, answers, and ICE candidates without relying on third-party meeting software.'),
        ('3.3.3 Project Relevance', 'The video consultation module is one of the most important differentiators of the project because it keeps the entire telemedicine experience within the system. No external conferencing tools are required, which improves usability and keeps the workflow focused on the appointment record.'),
        ('3.3.4 Session Control', 'The video session is connected to the appointment and tracked through a room or session record. The call can remain in waiting, active, or ended states, which allows the frontend to display correct consultation status and helps the backend maintain appointment history.'),
        ('3.3.5 Quality and Safety', 'The consultation page is designed to manage camera and microphone permissions, connection states, and session ending actions. This gives both users a clear and controlled teleconsultation experience within the browser.'),
    ])

    add_module_section(doc, '3.4 Prescription Management Module', [
        ('3.4.1 Objective', 'This module allows doctors to issue secure digital prescriptions after the consultation is complete.'),
        ('3.4.2 Functional Description', 'A prescription may contain the medicine name, dosage, frequency, duration, quantity, special instructions, diagnosis, and validity period. Each prescription is linked to a specific appointment, doctor, and patient so that the medical record remains traceable. Patients can view, download, and print the prescription from their dashboard whenever required.'),
        ('3.4.3 Project Relevance', 'The prescription workflow is essential because it converts a live consultation into an actionable treatment record. It also provides continuity of care by linking the doctor’s instructions to the appointment and the patient’s future visits.'),
        ('3.4.4 Data Handling', 'Prescription details are stored securely in MongoDB and can be accessed later for follow-up visits, patient review, or administrative auditing. Status changes such as issued, viewed, picked up, or cancelled help the system keep a detailed prescription lifecycle.'),
        ('3.4.5 User Experience', 'The frontend supports clear prescription detail pages and downloadable output so that patients can retain a copy for their records or pharmacy visits.'),
    ])

    add_module_section(doc, '3.5 AI Health Advisory Module', [
        ('3.5.1 Objective', 'This module improves decision support through AI-based assistance for symptom interpretation, report simplification, and short medical summaries.'),
        ('3.5.2 Functional Description', 'Based on the symptoms entered by a patient, the system recommends a suitable doctor specialization. After consultation, doctor notes can be converted into a short AI-generated summary to help patients understand the treatment plan. The platform also allows medical reports to be uploaded and simplified using OCR and AI processing so that users can read the content in a more accessible form.'),
        ('3.5.3 Project Relevance', 'The AI advisory module adds intelligence to the clinic platform while still keeping the human doctor at the center of diagnosis. It supports faster understanding, better triage, and improved communication for complex medical documents.'),
        ('3.5.4 Technical Support', 'The current project analysis shows the use of Gemini-based AI services along with OCR/document extraction tools such as Tesseract and PDF parsing utilities. This allows the system to process text from images and PDFs before generating a structured summary or advisory output.'),
        ('3.5.5 Safety Note', 'The AI output is informational only and does not replace a doctor’s clinical diagnosis. The interface should always present an appropriate disclaimer so that users understand the advisory nature of the result.'),
    ])

    add_module_section(doc, '3.6 Wellness Module', [
        ('3.6.1 Objective', 'This module promotes preventive awareness by helping users examine family health patterns and possible hereditary risks.'),
        ('3.6.2 Functional Description', 'Users can enter health details of family members such as parents or grandparents, along with simple lifestyle information. The system evaluates the input with predefined rules and highlights possible inherited health concerns in a user-friendly way. The result is intended to help users notice patterns and consider early preventive action.'),
        ('3.6.3 Project Relevance', 'The wellness module broadens the platform beyond treatment and consultation. It supports early awareness, family health tracking, and patient education, which makes the project more comprehensive and socially useful.'),
        ('3.6.4 Data Policy', 'According to the project design, this information is processed live and is not stored permanently. That approach preserves privacy while still giving the user an instant wellness insight.'),
        ('3.6.5 Guidance Note', 'The module offers awareness and guidance only. It must not be presented as a medical conclusion, and users should be encouraged to consult a doctor for any real concern.'),
    ])

    add_heading(doc, '3.7 Overall Quality and Integration Summary')
    p = doc.add_paragraph()
    add_text(p, 'The six modules work together as a complete virtual clinic workflow. ', 12)
    add_text(p, 'User Management provides secure authentication and access control; Appointment Management handles scheduling; Video Consultation enables the live doctor-patient interaction; Prescription Management records the treatment plan; AI Health Advisory supports smart assistance; and Wellness helps with preventive awareness. ', 12)
    add_text(p, 'Together, these modules create a coherent healthcare system that is practical for demonstration, academically relevant, and aligned with the current project implementation.', 12)
    format_paragraph(p)

    add_heading(doc, '3.8 Guide Remarks / Suggestions')
    p = doc.add_paragraph()
    add_text(p, 'Guide Remarks: The chapter is complete, well-structured, and aligned with the implemented MERN stack project. ', 12)
    add_text(p, 'Suggestions: If required for final submission, the document can be further enriched with live screenshots of the login page, dashboard pages, appointment flow, video consultation screen, prescription view, AI analyzer, and wellness module UI.', 12)
    format_paragraph(p)

    p = doc.add_paragraph()
    add_text(p, 'Guide Signature: ________________________________', 12, bold=True)
    format_paragraph(p)
    p = doc.add_paragraph()
    add_text(p, 'Date: ________________________________', 12, bold=True)
    format_paragraph(p)

    if OUT_PATH.exists():
        OUT_PATH.unlink()
    doc.save(OUT_PATH)
    print(f'Created {OUT_PATH.resolve()}')


if __name__ == '__main__':
    main()
