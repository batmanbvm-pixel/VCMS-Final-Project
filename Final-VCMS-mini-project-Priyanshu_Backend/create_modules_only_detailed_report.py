from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path('/Users/patelpreet/Desktop/VCMS_Modules_Only_Report_Detailed.docx')


def set_font(run, size=12, bold=False, italic=False, color=None):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def text(doc, value, size=12, bold=False, italic=False, color=None, align=None):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(value)
    set_font(r, size=size, bold=bold, italic=italic, color=color)
    return p


def heading(doc, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(value)
    set_font(r, size=14, bold=True, color='0B4F6C')
    return p


def subheading(doc, value):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(value)
    set_font(r, size=12, bold=True, color='1F2937')
    return p


def bullet(doc, value):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-0.35)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(value)
    set_font(r, size=12)


def add_module_block(doc, module_no, module_title, objective, submodules, new_page=False):
    if new_page:
        doc.add_page_break()

    heading(doc, f'{module_no}. {module_title}')
    obj = text(doc, f'Objective: {objective}')
    obj.paragraph_format.left_indent = Cm(0.5)
    obj.paragraph_format.space_after = Pt(6)
    for idx, sm in enumerate(submodules, start=1):
        subheading(doc, f'{module_no}.{idx} {sm["title"]}')
        intro = text(doc, sm['intro'])
        intro.paragraph_format.left_indent = Cm(0.8)
        intro.paragraph_format.space_after = Pt(3)
        for point in sm['points']:
            bullet(doc, point)
        gap = doc.add_paragraph()
        gap.paragraph_format.space_after = Pt(3)


def main():
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Cm(21)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(12)

    text(doc, 'VCMS Project Report (Modules Only)', size=18, bold=True, color='0B4F6C', align=WD_ALIGN_PARAGRAPH.CENTER)
    text(doc, 'Detailed Pointwise Version (Viva Ready)', size=15, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    text(doc, 'Virtual Clinic Management System', size=13, italic=True, align=WD_ALIGN_PARAGRAPH.CENTER)
    text(doc, 'Only six modules are included below. Architecture-level discussion is intentionally excluded.', size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_module_block(
        doc,
        '1',
        'User Management Module',
        'This module manages user identity, authentication, authorization, and account lifecycle for Patient, Doctor, and Admin users.',
        [
            {
                'title': 'Role Model and Access Governance',
                'intro': 'User governance is implemented as a strict role-based structure across pages, APIs, and data access.',
                'points': [
                    'Three core roles are supported: Patient, Doctor, and Admin.',
                    'Protected frontend routes ensure role-specific dashboards are isolated and secure.',
                    'Backend middleware validates user role before allowing access to restricted operations.',
                    'Admin role includes approval and status control for account governance workflows.',
                    'Notification, settings, and profile personalization flows are role-aware and tied to this module.'
                ]
            },
            {
                'title': 'Registration, Login, and Account Operations',
                'intro': 'Account creation and session flows cover the full user lifecycle from onboarding to secure re-login.',
                'points': [
                    'Registration fields include name, email, phone, password, and selected role.',
                    'Doctor profile fields include specialization, consultation fee, and professional data.',
                    'Patient profile supports age/date-of-birth and basic health background details.',
                    'Login, logout, profile update, password reset, and token refresh workflows are integrated.',
                    'Session continuity supports protected access to appointments, prescriptions, AI analyzer, and video pages.'
                ]
            },
            {
                'title': 'Security and Persistence',
                'intro': 'Security controls are integrated at credential, session, and API levels.',
                'points': [
                    'Passwords are stored as bcrypt hashes, never plain text.',
                    'JWT tokens secure authenticated API usage and protected session flow.',
                    'Input validation/sanitization guards are applied on critical auth routes.',
                    'Users collection stores role metadata, profile data, approval flags, and account status fields.',
                    'Audit-friendly account metadata supports admin governance and compliance discussions in viva.'
                ]
            }
        ]
    )

    add_module_block(
        doc,
        '2',
        'Appointment Management Module',
        'This module provides complete appointment scheduling and tracking from booking request to consultation completion.',
        [
            {
                'title': 'Booking Methods and User Flow',
                'intro': 'Appointments can be created through manual and guided interaction modes for better usability.',
                'points': [
                    'Manual flow allows patient to select doctor, date, and time slot directly.',
                    'Chatbot flow asks symptoms, specialization preference, date, and time to simplify booking.',
                    'Doctor and admin dashboards expose role-specific appointment views and actions.',
                    'Public doctor discovery and specialization filtering feed this booking pipeline.'
                ]
            },
            {
                'title': 'Status Lifecycle and Business Rules',
                'intro': 'Each booking follows a controlled status model backed by validation rules.',
                'points': [
                    'Status pipeline includes Booked, Accepted, In Progress, Completed, Rejected, and Cancelled.',
                    'Slot conflict checks and schedule validation prevent invalid or duplicate bookings.',
                    'Reason-based rejection/cancellation support keeps decisions auditable.',
                    'Role-specific actions: patient book/cancel, doctor accept/reject, admin supervise globally.',
                    'Lifecycle transitions are reflected in real time to maintain cross-dashboard consistency.'
                ]
            },
            {
                'title': 'Realtime Integration and Data Model',
                'intro': 'Appointment data drives downstream modules and real-time user awareness.',
                'points': [
                    'Socket-based updates push immediate appointment status changes to active users.',
                    'Appointment APIs support create/list/get/update/cancel/accept/reject/slot checking.',
                    'Appointments collection stores patientId, doctorId, symptoms, notes, status, and timestamps.',
                    'Confirmed appointments act as trigger point for video consultation and prescription flow.',
                    'Medical-history context and notification records are linked to appointment events.'
                ]
            }
        ],
        new_page=True
    )

    add_module_block(
        doc,
        '3',
        'Video Consultation Module',
        'This module enables secure in-app teleconsultation so doctor and patient can consult in real time without external meeting tools.',
        [
            {
                'title': 'Consultation Session Flow',
                'intro': 'Video sessions are linked with appointment context to preserve clinical continuity.',
                'points': [
                    'Doctor and patient join consultation room mapped to scheduled appointment context.',
                    'Session lifecycle states include waiting, active, and ended.',
                    'In-app call flow avoids context switching to third-party platforms.',
                    'Consultation continuity is preserved because patient, doctor, and appointment records remain in one system.'
                ]
            },
            {
                'title': 'WebRTC and Signaling',
                'intro': 'Media and signaling layers are implemented using modern web real-time standards.',
                'points': [
                    'WebRTC handles peer audio/video connection.',
                    'Socket.IO signaling exchanges offer, answer, ICE candidate, and call events.',
                    'Connection-state events support resilient session setup and user feedback.',
                    'Realtime chat and notification hooks support communication before and during consultation.'
                ]
            },
            {
                'title': 'UI Controls and Persistence',
                'intro': 'User interface and backend routes maintain full consultation traceability.',
                'points': [
                    'Video UI supports mic/camera controls, join/leave, and connection indicators.',
                    'Video APIs cover room create/fetch/status updates.',
                    'VideoSessions collection stores appointment linkage, participants, roomId, and timestamps.',
                    'Call completion feeds downstream prescription and appointment-finalization actions.'
                ]
            }
        ],
        new_page=True
    )

    add_module_block(
        doc,
        '4',
        'Prescription Management Module',
        'This module captures treatment outcomes as digital prescriptions linked to patient consultation records.',
        [
            {
                'title': 'Prescription Lifecycle and Usage',
                'intro': 'Prescription records are generated after consultation and remain available for follow-up care.',
                'points': [
                    'Doctors can create and issue prescriptions after consultation completion.',
                    'Patients can access active and historical prescriptions from dashboard interfaces.',
                    'Status-based actions support prescription tracking across issue/view/update events.',
                    'Prescription workflow is integrated with appointment closure and doctor follow-up flow.'
                ]
            },
            {
                'title': 'Clinical Data Structure',
                'intro': 'Prescription schema is designed for practical medication management and traceability.',
                'points': [
                    'Data includes medicine name, dosage, frequency, duration, quantity, and instructions.',
                    'Diagnosis notes and validity metadata support contextual treatment continuity.',
                    'Each record is linked to doctorId, patientId, and appointmentId references for traceability.',
                    'Medical history alignment allows treatment progression to be explained in viva demonstrations.'
                ]
            },
            {
                'title': 'Frontend, Backend, and Database Coverage',
                'intro': 'Prescription functionality is fully integrated across views, APIs, and storage.',
                'points': [
                    'Frontend includes create form, list pages, detail views, and patient-readable output.',
                    'Backend supports create/list/get/update/issue and appointment-linked retrieval.',
                    'Prescriptions collection stores medicine arrays and lifecycle metadata for future audit/use.',
                    'PDF/view flow and notification triggers improve practical usability for patients.'
                ]
            }
        ],
        new_page=True
    )

    add_module_block(
        doc,
        '5',
        'AI Health Advisory Module',
        'This module provides AI-assisted guidance for symptom routing, report understanding, and patient-friendly clinical summarization.',
        [
            {
                'title': 'Decision Support and Summarization',
                'intro': 'AI features improve patient understanding while preserving doctor authority.',
                'points': [
                    'Symptom-driven specialization recommendations guide patients to suitable doctors.',
                    'Doctor notes can be transformed into concise, user-readable summaries.',
                    'Advisory outputs improve triage clarity and reduce confusion for non-technical users.',
                    'AI guidance is designed to support, not replace, clinical decision by doctors.'
                ]
            },
            {
                'title': 'OCR and Document Simplification',
                'intro': 'Complex medical documents are converted into readable output using extraction + AI processing.',
                'points': [
                    'Uploaded images/PDFs are parsed through OCR/text extraction pipeline.',
                    'Extracted content is summarized into structured, simplified medical guidance.',
                    'Result sections are presented in easy-to-understand blocks on frontend analyzer pages.',
                    'Multi-step analysis improves readability of complex pathology and clinical documents.'
                ]
            },
            {
                'title': 'Technical and Safety Controls',
                'intro': 'AI integration includes both functionality and medical safety boundaries.',
                'points': [
                    'Frontend includes upload, progress, and summary rendering workflows.',
                    'Backend AI endpoints support extraction, summarization, and formatted response generation.',
                    'System explicitly marks outputs as informational only, not final diagnosis.',
                    'Error handling and fallback behavior improve reliability during viva demos.'
                ]
            }
        ],
        new_page=True
    )

    add_module_block(
        doc,
        '6',
        'Wellness Module',
        'This module focuses on preventive awareness by analyzing family health indicators and risk tendencies.',
        [
            {
                'title': 'Preventive Objective and Scope',
                'intro': 'Wellness logic extends platform value beyond treatment toward early awareness.',
                'points': [
                    'Users submit family-health indicators and related lifestyle context.',
                    'Module helps identify potential hereditary awareness signals early.',
                    'Output complements consultation decisions but does not replace them.',
                    'Module extends project scope from treatment-only to preventive-care orientation.'
                ]
            },
            {
                'title': 'Risk Processing and Output Design',
                'intro': 'Risk awareness output is intentionally simple, readable, and non-clinical.',
                'points': [
                    'Rule-based live processing generates awareness-oriented feedback.',
                    'Result style is low/moderate/high guidance rather than medical diagnosis.',
                    'Users are directed to seek doctor consultation for clinical conclusions.',
                    'Presentation is intentionally simple so non-technical users can understand risk tendencies quickly.'
                ]
            },
            {
                'title': 'Privacy and Frontend Experience',
                'intro': 'Wellness UX is designed for quick awareness with strong privacy posture.',
                'points': [
                    'Data is processed live and not stored as permanent clinical record.',
                    'Frontend supports guided questionnaire/dialog with immediate result view.',
                    'Design increases health awareness while minimizing persistence risks.',
                    'Privacy-first behavior is aligned with healthcare sensitivity expectations.'
                ]
            }
        ],
        new_page=True
    )

    heading(doc, 'Guide Remarks / Suggestions')
    text(doc, 'Guide Remark: Detailed module-only report prepared with complete project feature mapping.', bold=True)
    text(doc, 'Guide Signature: ______________________________', bold=True)
    text(doc, 'Date: ______________________________', bold=True)

    if OUT.exists():
        OUT.unlink()
    doc.save(OUT)
    print(f'CREATED: {OUT}')


if __name__ == '__main__':
    main()
