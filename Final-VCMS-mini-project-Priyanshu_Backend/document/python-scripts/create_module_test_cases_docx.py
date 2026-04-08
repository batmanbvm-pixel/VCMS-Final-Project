from docx import Document
from docx.shared import Pt, Cm
from docx.oxml.ns import qn
from pathlib import Path

OUT = Path('/Users/patelpreet/Desktop/VCMS_Module_Test_Cases_Main.docx')


def set_run_font(run, size=12, bold=False):
    run.font.name = 'Times New Roman'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    run.font.size = Pt(size)
    run.font.bold = bold


def add_text_paragraph(doc, text, size=12, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)
    return p


def style_cell_text(cell, text, bold=False, size=11):
    cell.text = ''
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    set_run_font(r, size=size, bold=bold)


def add_module_table(doc, module_title, rows):
    add_text_paragraph(doc, module_title, size=13, bold=True)

    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'

    headers = ['Test Case', 'Test Data', 'Test Result', 'Test Report']
    for i, h in enumerate(headers):
        style_cell_text(table.rows[0].cells[i], h, bold=True, size=12)

    for row in rows:
        tr = table.add_row().cells
        style_cell_text(tr[0], row[0])
        style_cell_text(tr[1], row[1])
        style_cell_text(tr[2], row[2])
        style_cell_text(tr[3], row[3])

    # Adjust widths similar to screenshot layout
    for r in table.rows:
        r.cells[0].width = Cm(5.0)
        r.cells[1].width = Cm(7.5)
        r.cells[2].width = Cm(2.5)
        r.cells[3].width = Cm(4.0)

    doc.add_paragraph('')


def main():
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(2.2)
    section.right_margin = Cm(2.2)

    normal = doc.styles['Normal']
    normal.font.name = 'Times New Roman'
    normal._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    normal.font.size = Pt(11)

    add_text_paragraph(doc, 'VCMS Project - Main Module Test Cases', size=16, bold=True)
    add_text_paragraph(doc, 'Reference Format: Test Case | Test Data | Test Result | Test Report', size=11, bold=False)
    add_text_paragraph(doc, 'Note: Main test scenarios only (concise version for viva/reference).', size=11, bold=False)
    doc.add_paragraph('')

    add_module_table(doc, '1. User Management Module', [
        ('Blank registration fields', 'Name/Email/Phone/Password empty', 'Invalid', 'Validation message shown'),
        ('Existing email registration', 'Email: existing user, valid password', 'Invalid', 'Email already registered'),
        ('Valid patient registration', 'New email + valid patient details', 'Valid', 'Account created successfully'),
        ('Valid doctor registration', 'New email + specialization + fee', 'Valid', 'Doctor account created'),
        ('Valid login', 'Registered email + correct password', 'Valid', 'Dashboard opens by role'),
    ])

    add_module_table(doc, '2. Appointment Management Module', [
        ('Book without slot selection', 'Doctor selected, date missing time', 'Invalid', 'Select date/time prompt'),
        ('Manual appointment booking', 'Patient + doctor + date + slot', 'Valid', 'Appointment created with pending status'),
        ('Chatbot guidance response', 'Chat message: appointment help', 'Valid', 'Bot returns booking guidance steps'),
        ('Double-book same slot', 'Same doctor/date/time already booked', 'Invalid', 'Slot conflict message'),
        ('Doctor accepts appointment', 'Pending appointment action = Accept', 'Valid', 'Status updated to confirmed'),
    ])

    add_module_table(doc, '3. Video Consultation Module', [
        ('Join without valid room', 'Random/invalid room ID', 'Invalid', 'Room/session not found'),
        ('Patient joins scheduled room', 'Valid appointment room ID', 'Valid', 'Connected to consultation room'),
        ('Doctor joins same room', 'Matching room ID and role', 'Valid', 'Video session active'),
        ('WebRTC signaling exchange', 'Offer/answer/ICE events', 'Valid', 'Stable media connection'),
        ('End consultation', 'Doctor/patient clicks end call', 'Valid', 'Session closed and logged'),
    ])

    add_module_table(doc, '4. Prescription Management Module', [
        ('Create with missing medicine', 'Dosage entered, medicine empty', 'Invalid', 'Medicine required message'),
        ('Create valid prescription', 'Medicine + dosage + duration + instructions', 'Valid', 'Prescription saved'),
        ('Link to appointment check', 'Valid appointmentId provided', 'Valid', 'Prescription linked correctly'),
        ('Patient view prescription', 'Patient opens prescription tab', 'Valid', 'Current + history shown'),
        ('Issue draft prescription', 'Doctor issues existing draft prescription', 'Valid', 'Status changes draft → issued'),
    ])

    add_module_table(doc, '5. AI Health Advisory Module', [
        ('Prescription summary generation', 'POST /ai/summarize with type + content', 'Valid', 'Summary and key points returned'),
        ('Missing summarize payload', 'POST /ai/summarize without content/type', 'Invalid', '400 Missing content or type'),
        ('Medical report analysis', 'POST /ai/analyze-report with report text', 'Valid', 'Report summary and recommendations returned'),
        ('Unreadable report content', 'POST /ai/analyze-report with very short content', 'Invalid', '422 readable-content error returned'),
        ('OCR without file upload', 'POST /ai/extract-text without file', 'Invalid', '400 No file provided'),
    ])

    add_module_table(doc, '6. Wellness Module', [
        ('Open Family Risk dialog', 'Patient dashboard → Family Risk card', 'Valid', 'Risk predictor dialog opens'),
        ('Run inherited risk predictor', 'Select disease + answer predefined questions', 'Valid', 'Risk % and risk level shown'),
        ('No disease selected', 'Open predictor with no disease chosen', 'Invalid', 'Prompt to select disease first'),
        ('Reset risk answers', 'Use reset action in Family Risk dialog', 'Valid', 'Answers return to default state'),
        ('Guidance disclaimer check', 'View risk summary text in dialog', 'Valid', 'Advisory-only wording displayed'),
    ])

    doc.save(OUT)
    print(f'CREATED: {OUT}')


if __name__ == '__main__':
    main()
